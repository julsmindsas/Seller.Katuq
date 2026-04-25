from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def style_doc(doc):
    """Configura estilos base del documento"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            h.font.size = Pt(20)
            h.font.bold = True
        elif level == 2:
            h.font.size = Pt(16)
            h.font.bold = True
        elif level == 3:
            h.font.size = Pt(13)
            h.font.bold = True

def add_table(doc, headers, rows, col_widths=None):
    """Agrega una tabla formateada"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table

def add_code_block(doc, code):
    """Agrega un bloque de codigo"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(1)

# ============================================================
# DOCUMENTO 1: DOCUMENTACION TECNICA
# ============================================================
def generate_technical_doc():
    doc = Document()
    style_doc(doc)

    # Portada
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('Sistema de Notificaciones Katuq', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Documentacion Tecnica')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Version 1.0 | Marzo 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # Tabla de contenidos
    doc.add_heading('Tabla de Contenidos', level=1)
    toc_items = [
        '1. Arquitectura General',
        '2. Flujos de Notificacion',
        '3. Estados que Generan Notificaciones',
        '4. Preferencias por Empresa',
        '5. Canales de Envio',
        '6. URL Corta para SMS',
        '7. Estructura de Archivos',
        '8. Configuracion del Entorno (.env)',
        '9. Seguridad',
        '10. Dependencias Externas',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # 1. Arquitectura General
    doc.add_heading('1. Arquitectura General', level=1)

    doc.add_heading('1.1 Descripcion', level=2)
    doc.add_paragraph(
        'El sistema de notificaciones de Katuq permite enviar notificaciones multicanal '
        '(Email, SMS) a los clientes cuando cambian los estados de sus pedidos. Cada empresa '
        'puede configurar individualmente que canales y estados activan notificaciones.'
    )

    doc.add_heading('1.2 Componentes Principales', level=2)
    add_table(doc,
        ['Componente', 'Tecnologia', 'Ubicacion'],
        [
            ['Panel de Notificaciones', 'Angular 14', 'src/app/components/notificaciones/'],
            ['API de Preferencias', 'Express/Node.js', 'controllers/notificationPreferences.js'],
            ['Envio Directo (Email)', 'Nodemailer SMTP', 'controllers/orders.js'],
            ['Envio Directo (SMS)', 'labsMobile API', 'services/smsService.js'],
            ['Templates Email', 'HTML modular', 'services/notifications/templates/'],
            ['Templates SMS', 'Texto con variables', 'services/notifications/smsTemplates.js'],
            ['Email Preview', 'Firebase Storage', 'controllers/orders.js (saveEmailPreview)'],
            ['URL Corta', 'Express redirect', 'index.js (GET /v1/e/:id)'],
            ['Preferencias Storage', 'Firestore', 'company_notification_preferences'],
        ]
    )

    doc.add_heading('1.3 Repositorios', level=2)
    add_table(doc,
        ['Repositorio', 'Rama', 'Descripcion'],
        [
            ['Seller.Katuq', 'feature/merge-Notifications', 'Frontend Angular - Panel de preferencias'],
            ['katuq_admin_back_firebase', 'backend-notifications', 'Backend Express - API, envio, templates'],
        ]
    )

    # 2. Flujos
    doc.add_heading('2. Flujos de Notificacion', level=1)

    doc.add_heading('2.1 Pedido Editado (cambio de estado)', level=2)
    steps = [
        'Frontend llama PUT /v1/orders/:id con el nuevo estado',
        'controllers/orders.js > edit() procesa la actualizacion',
        'setImmediate (no bloquea la respuesta HTTP):',
        '   a. sendDirectTemplateEmail() genera HTML, envia email, guarda preview en Storage',
        '   b. sendDirectSms() usa el link del preview como linkPedido en el SMS',
        'El SMS llega al cliente con URL corta: https://back.katuq.com/v1/e/{docId}',
        'Cliente abre el link > redirect a Storage > ve el email completo en el navegador',
    ]
    for i, step in enumerate(steps):
        doc.add_paragraph(step, style='List Number' if not step.startswith('   ') else 'List Bullet')

    doc.add_heading('2.2 Pedido Creado', level=2)
    steps = [
        'Frontend llama POST /v1/orders con los datos del pedido',
        'controllers/orders.js > create() guarda el pedido',
        'El email legacy del frontend (con carrito, productos) se envia normalmente',
        'setImmediate:',
        '   a. notificationHooks.onOrderCreated() encola en Firestore',
        '   b. saveEmailPreview(emailHtml) guarda el HTML en Storage',
        '   c. sendDirectCreatedSms(order, previewUrl) envia SMS con link al preview',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number' if not step.startswith('   ') else 'List Bullet')

    doc.add_heading('2.3 Pedido Editado por Transportador', level=2)
    doc.add_paragraph(
        'Mismo flujo que "Pedido Editado" pero a traves del endpoint editByTransporter(). '
        'El transportador cambia el estado (ej: Entregado, Rechazado) y se disparan las '
        'mismas notificaciones de email + SMS.'
    )

    # 3. Estados
    doc.add_heading('3. Estados que Generan Notificaciones', level=1)

    doc.add_heading('3.1 Estados de Proceso (estadoProceso)', level=2)
    add_table(doc,
        ['Estado', 'Notifica al Cliente', 'Template Email', 'Template SMS'],
        [
            ['SinProducir', 'No', '-', '-'],
            ['EnProduccion', 'No', '-', '-'],
            ['ProducidoParcialmente', 'No', '-', '-'],
            ['ProducidoTotalmente', 'Si', 'ORDER_ProducidoTotalmente', 'ORDER_ProducidoTotalmente'],
            ['ParaDespachar', 'No', '-', '-'],
            ['Empacado', 'No', '-', '-'],
            ['Despachado', 'Si', 'ORDER_Despachado', 'ORDER_Despachado'],
            ['Entregado', 'Si', 'ORDER_Entregado', 'ORDER_Entregado'],
            ['Rechazado', 'Si', 'ORDER_Rechazado', 'ORDER_Rechazado'],
        ]
    )

    doc.add_heading('3.2 Estados de Pago (estadoPago)', level=2)
    add_table(doc,
        ['Estado', 'Notifica al Cliente', 'Template Email', 'Template SMS'],
        [
            ['Pendiente', 'No', '-', '-'],
            ['PreAprobado', 'No', '-', '-'],
            ['Aprobado', 'Si', 'PAYMENT_Aprobado', 'PAYMENT_Aprobado'],
            ['Rechazado', 'No', '-', '-'],
            ['Cancelado', 'No', '-', '-'],
        ]
    )

    doc.add_heading('3.3 Creacion de Pedido', level=2)
    add_table(doc,
        ['Evento', 'Template Email', 'Template SMS'],
        [
            ['Pedido creado', 'Email legacy (frontend)', 'ORDER_CREATED'],
        ]
    )

    # 4. Preferencias
    doc.add_heading('4. Preferencias por Empresa', level=1)

    doc.add_heading('4.1 Estructura en Firestore', level=2)
    doc.add_paragraph('Collection: company_notification_preferences')
    doc.add_paragraph('Document ID: nombre de la empresa (ej: "Tienda Demo KAI Import")')
    add_code_block(doc, '''
{
  "notifications": {
    "order_created": true,
    "payment_approved": false,
    "order_produced": true,
    "order_dispatched": true,
    "order_delivered": true,
    "order_rejected": false
  },
  "sms_notifications": {
    "order_created": false,
    "payment_approved": false,
    "order_produced": false,
    "order_dispatched": true,
    "order_delivered": true,
    "order_rejected": false
  },
  "company": "Tienda Demo KAI Import",
  "updatedAt": "2026-03-27T..."
}''')

    doc.add_heading('4.2 Endpoints API', level=2)
    add_table(doc,
        ['Metodo', 'Ruta', 'Descripcion'],
        [
            ['GET', '/v1/notification-preferences/company/:companyName', 'Obtener preferencias'],
            ['PUT', '/v1/notification-preferences/company/:companyName', 'Guardar preferencias'],
        ]
    )

    doc.add_paragraph(
        'Logica de defaults: Todas las preferencias son false por defecto (opt-in). '
        'Cada empresa activa individualmente los canales y estados que desea.'
    )

    # 5. Canales
    doc.add_heading('5. Canales de Envio', level=1)

    doc.add_heading('5.1 Email', level=2)
    add_table(doc,
        ['Caracteristica', 'Detalle'],
        [
            ['Proveedor', 'Nodemailer con SMTP'],
            ['Templates', 'HTML modular en services/notifications/templates/ (14 archivos)'],
            ['Preview', 'HTML guardado en Firebase Storage como archivo publico'],
            ['Branding', 'Logo y colores por empresa via brandingService.js'],
        ]
    )

    doc.add_heading('5.2 SMS', level=2)
    add_table(doc,
        ['Caracteristica', 'Detalle'],
        [
            ['Proveedor', 'labsMobile (API REST)'],
            ['Endpoint API', 'https://api.labsmobile.com/json/send'],
            ['Autenticacion', 'Basic Auth (usuario + API key)'],
            ['Servicio', 'services/smsService.js'],
            ['Templates', 'services/notifications/smsTemplates.js'],
            ['Limite', '160 caracteres por mensaje'],
        ]
    )

    doc.add_heading('5.2.1 Templates SMS', level=3)
    add_table(doc,
        ['Tipo', 'Mensaje'],
        [
            ['ORDER_CREATED', '{empresa}: Tu pedido #{nro} ha sido confirmado por {total}. Detalles: {link}'],
            ['PAYMENT_Aprobado', '{empresa}: El pago de tu pedido #{nro} ha sido aprobado. Detalles: {link}'],
            ['ORDER_ProducidoTotalmente', '{empresa}: Tu pedido #{nro} ha sido producido y esta listo para despacho. Detalles: {link}'],
            ['ORDER_Despachado', '{empresa}: Tu pedido #{nro} esta en camino. Guia: {guia}. Detalles: {link}'],
            ['ORDER_Entregado', '{empresa}: Tu pedido #{nro} ha sido entregado. Gracias por tu compra! Detalles: {link}'],
            ['ORDER_Rechazado', '{empresa}: Tu pedido #{nro} no pudo ser procesado. Detalles: {link}'],
        ]
    )

    doc.add_heading('5.3 WhatsApp', level=2)
    doc.add_paragraph('Estado: Deshabilitado (proximamente). Columna decorativa en el panel de notificaciones.')

    # 6. URL Corta
    doc.add_heading('6. URL Corta para SMS', level=1)

    doc.add_heading('6.1 Problema', level=2)
    doc.add_paragraph(
        'Las URLs de Firebase Storage son muy largas (~120 caracteres), dejando poco espacio '
        'para el mensaje SMS (limite 160 caracteres).'
    )

    doc.add_heading('6.2 Solucion', level=2)
    add_table(doc,
        ['Tipo', 'URL', 'Caracteres'],
        [
            ['URL larga (Storage)', 'https://storage.googleapis.com/julsmind-katuq.appspot.com/email_previews/...', '~120'],
            ['URL corta (backend)', 'https://back.katuq.com/v1/e/abc123xyz', '~42'],
        ]
    )

    doc.add_heading('6.3 Flujo Tecnico', level=2)
    steps = [
        'saveEmailPreview() guarda HTML en Storage + metadata en Firestore (collection email_previews)',
        'Captura el docRef.id del documento de Firestore',
        'Si BACKEND_PUBLIC_URL existe, retorna {BACKEND_PUBLIC_URL}/v1/e/{docId}',
        'Si no existe (desarrollo local), retorna la URL de Storage directa',
        'GET /v1/e/:id busca el documento y hace redirect 301 a la URL de Storage',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # 7. Estructura de Archivos
    doc.add_heading('7. Estructura de Archivos', level=1)

    doc.add_heading('7.1 Frontend (Seller.Katuq)', level=2)
    add_table(doc,
        ['Archivo', 'Descripcion'],
        [
            ['components/notificaciones/notificaciones.component.ts', 'Componente principal del panel'],
            ['components/notificaciones/notificaciones.component.html', 'Template con toggles Email/SMS/WhatsApp'],
            ['components/notificaciones/notificaciones.component.scss', 'Estilos del panel'],
            ['components/notificaciones/notificaciones.module.ts', 'Modulo Angular'],
            ['components/notificaciones/notificaciones-routing.module.ts', 'Routing lazy-loaded'],
            ['shared/routes/routes.ts', 'Ruta /notificaciones con AuthGuard'],
            ['shared/services/maestros/maestro.service.ts', 'get/saveCompanyNotificationPreferences()'],
            ['shared/services/notifications/notification.types.ts', 'Enums y tipos'],
        ]
    )

    doc.add_heading('7.2 Backend (katuq_admin_back_firebase)', level=2)
    add_table(doc,
        ['Archivo', 'Descripcion'],
        [
            ['controllers/orders.js', 'create(), edit(), sendDirectSms(), saveEmailPreview()'],
            ['controllers/notificationPreferences.js', 'CRUD preferencias empresa/usuario'],
            ['routers/notifications.js', '/pause, /resume, /status, /preview'],
            ['routers/notificationsSend.js', 'POST /v1/notifications/send'],
            ['services/smsService.js', 'labsMobile API (envio SMS)'],
            ['services/brandingService.js', 'Logo/colores por empresa'],
            ['services/notifications/smsTemplates.js', 'Templates SMS con variables'],
            ['services/notifications/templateHelpers.js', 'Render HTML, bloques opcionales'],
            ['services/notifications/notificationHooks.js', 'Deteccion de cambios de estado'],
            ['services/notifications/notificationQueue.js', 'Cola de notificaciones (Firestore)'],
            ['services/notifications/templates/', '14 templates modulares HTML'],
            ['index.js', 'Registro de rutas Express + /v1/e/:id'],
        ]
    )

    # 8. Configuracion
    doc.add_heading('8. Configuracion del Entorno (.env)', level=1)

    doc.add_heading('8.1 Variables de Notificaciones', level=2)
    add_table(doc,
        ['Variable', 'Valor', 'Descripcion'],
        [
            ['ENABLE_ORDER_NOTIFICATIONS', 'true', 'Habilita el sistema de notificaciones'],
            ['ENABLE_EMAIL_NOTIFICATIONS', 'true', 'Habilita envio de emails'],
            ['ENABLE_SMS_NOTIFICATIONS', 'true', 'Habilita envio de SMS'],
            ['ENABLE_WHATSAPP_NOTIFICATIONS', 'false', 'WhatsApp deshabilitado'],
        ]
    )

    doc.add_heading('8.2 Filtro de Empresas', level=2)
    add_table(doc,
        ['Variable', 'Valor', 'Descripcion'],
        [
            ['ALLOWED_NOTIFICATION_COMPANIES', 'Tienda Demo KAI Import', 'Solo estas empresas envian notificaciones (vacio = todas)'],
        ]
    )

    doc.add_heading('8.3 labsMobile (SMS)', level=2)
    add_table(doc,
        ['Variable', 'Descripcion'],
        [
            ['LABSMOBILE_USERNAME', 'Usuario API labsMobile'],
            ['LABSMOBILE_API_KEY', 'API Key labsMobile'],
            ['LABSMOBILE_SENDER', 'Nombre remitente del SMS (default: Katuq)'],
        ]
    )

    doc.add_heading('8.4 URL Corta', level=2)
    add_table(doc,
        ['Variable', 'Valor', 'Descripcion'],
        [
            ['BACKEND_PUBLIC_URL', 'https://back.katuq.com', 'Base URL para URLs cortas en SMS'],
        ]
    )

    doc.add_heading('8.5 Modo Test vs Produccion', level=2)
    add_table(doc,
        ['Variable', 'Desarrollo', 'Produccion', 'Descripcion'],
        [
            ['NOTIFICATIONS_PUBLIC_MODE', 'false', 'true', 'false = todo va al test, true = va al cliente real'],
            ['NOTIFICATIONS_TEST_EMAIL', 'santygarciamartinez03@gmail.com', '-', 'Email de pruebas'],
            ['NOTIFICATIONS_TEST_PHONE', '573024218994', '-', 'Telefono de pruebas'],
        ]
    )

    doc.add_heading('8.6 Smart Notification Rules', level=2)
    add_table(doc,
        ['Variable', 'Valor', 'Descripcion'],
        [
            ['CUSTOMER_NOTIFY_STATES', 'ProducidoTotalmente,Despachado,Entregado,Rechazado', 'Estados de proceso que notifican'],
            ['CUSTOMER_NOTIFY_PAYMENT_STATES', 'Aprobado', 'Estados de pago que notifican'],
        ]
    )

    # 9. Seguridad
    doc.add_heading('9. Seguridad', level=1)
    items = [
        'Modo test: En desarrollo, TODOS los emails/SMS se redirigen a credenciales de test',
        'Filtro de empresas: ALLOWED_NOTIFICATION_COMPANIES limita que empresas pueden enviar',
        'Preferencias opt-in: Todo deshabilitado por defecto, cada empresa activa lo que necesita',
        'URL corta sin auth: /v1/e/:id es publico (necesario para links en SMS)',
        'Credenciales en .env: No se commitean al repositorio (gitignore)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # 10. Dependencias
    doc.add_heading('10. Dependencias Externas', level=1)
    add_table(doc,
        ['Servicio', 'Uso', 'Configuracion'],
        [
            ['Firebase Storage', 'Almacena HTML de emails publicos', 'Bucket: julsmind-katuq.appspot.com'],
            ['Firebase Firestore', 'Preferencias, metadata, cola', 'Collections: company_notification_preferences, email_previews'],
            ['labsMobile', 'Envio de SMS', 'API REST con Basic Auth'],
            ['AWS SQS', 'Cola de mensajes (legacy)', 'Queue: katuq-notifications'],
            ['Nodemailer', 'Envio de emails SMTP', 'Configurado en .env'],
        ]
    )

    filepath = os.path.join(OUTPUT_DIR, 'NOTIFICACIONES_DOC_TECNICA.docx')
    doc.save(filepath)
    print(f'Documento tecnico generado: {filepath}')


# ============================================================
# DOCUMENTO 2: DETALLE DE MERGE
# ============================================================
def generate_merge_doc():
    doc = Document()
    style_doc(doc)

    # Portada
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('Documento de Merge', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Sistema de Notificaciones - Detalle de Cambios')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Marzo 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # Info del merge
    doc.add_heading('Informacion del Merge', level=1)
    add_table(doc,
        ['Campo', 'Valor'],
        [
            ['Fecha', '2026-03-27'],
            ['Rama origen (frontend)', 'Notification'],
            ['Rama destino (frontend)', 'feature/merge-Notifications'],
            ['Rama (backend)', 'backend-notifications'],
            ['Commits frontend', '6 commits: 6bacc02e a edab8198'],
            ['Commits backend', '1 commit: a46c21a (Notification_SMS_Back_1)'],
        ]
    )

    # PARTE 1: FRONTEND
    doc.add_heading('PARTE 1: Frontend (Seller.Katuq)', level=1)
    doc.add_paragraph('Resumen: 22 archivos, +765 lineas, -1009 lineas')

    doc.add_heading('Commits Incluidos', level=2)
    add_table(doc,
        ['Commit', 'Mensaje', 'Descripcion'],
        [
            ['6bacc02e', 'Notificaciones Front 1.0', 'Componente base de notificaciones'],
            ['f2e88a45', 'Notificaciones Front 1.1', 'Mejoras UI'],
            ['c8688470', 'Notificaciones Front 1.2', 'Integracion con backend'],
            ['b2fb316b', 'Notificaciones-1.2', 'Ajustes preferencias'],
            ['4fc5bd90', 'Notification_email_1', 'Email funcional, SMS deshabilitado'],
            ['edab8198', 'Notification_SMS_Front_1', 'SMS habilitado en UI'],
        ]
    )

    # Archivos nuevos
    doc.add_heading('1. Archivos Nuevos', level=2)

    doc.add_heading('notificaciones.component.ts (+186 lineas)', level=3)
    doc.add_paragraph('Ruta: src/app/components/notificaciones/notificaciones.component.ts')
    add_table(doc,
        ['Lineas', 'Contenido'],
        [
            ['1-18', 'Imports: Component, OnInit, OnDestroy, takeUntil, Subject, ToastrService, MaestroService'],
            ['19-25', 'Interface NotificationPreferenceView: id, label, description, icon, channels (email, sms, whatsapp)'],
            ['26-45', 'Decorador @Component y variables: preferences, empresaActual, isLoading, isSaving, destroy$'],
            ['46-90', 'ngOnInit() - Inicializa 6 categorias: order_created, payment_approved, order_produced, order_dispatched, order_delivered, order_rejected'],
            ['91-137', 'loadPreferences() - Carga empresa de localStorage, GET al backend, aplica valores de notifications y sms_notifications'],
            ['139-155', 'toggleEmail(id) y toggleSms(id) - Invierte valor del canal y llama saveToFirestore()'],
            ['157-184', 'saveToFirestore() - Construye objetos, PUT al backend, toast exito/error'],
            ['185-186', 'ngOnDestroy() - Limpia subscripciones'],
        ]
    )

    doc.add_heading('notificaciones.component.html (+110 lineas)', level=3)
    doc.add_paragraph('Ruta: src/app/components/notificaciones/notificaciones.component.html')
    add_table(doc,
        ['Lineas', 'Contenido'],
        [
            ['1-20', 'Header con titulo "Preferencias de Notificaciones"'],
            ['21-50', 'Cabecera tabla: Tipo de Notificacion, Email, SMS (activo), WhatsApp (decorativo)'],
            ['51-95', '*ngFor sobre preferences: checkbox Email, checkbox SMS, checkbox WhatsApp deshabilitado'],
            ['96-110', 'Footer con boton guardar y spinner de carga'],
        ]
    )

    doc.add_heading('notificaciones.component.scss (+201 lineas)', level=3)
    doc.add_paragraph('Ruta: src/app/components/notificaciones/notificaciones.component.scss')
    doc.add_paragraph('Estilos del contenedor, tabla responsive, custom checkboxes con animaciones, columnas por canal, clases decorativas.')

    doc.add_heading('notificaciones.module.ts (+21 lineas)', level=3)
    doc.add_paragraph('Imports: CommonModule, TranslateModule, NotificacionesRoutingModule. Declarations: NotificacionesComponent.')

    doc.add_heading('notificaciones-routing.module.ts (+11 lineas)', level=3)
    doc.add_paragraph('Routing lazy-loaded: path "" > NotificacionesComponent.')

    doc.add_heading('notifications-send.route.js (+59 lineas)', level=3)
    doc.add_paragraph('Ruta: src/app/shared/routes/notifications-send.route.js')
    doc.add_paragraph('Utilidad para envio de notificaciones desde el frontend: buildNotificationPayload(), sendNotification().')

    doc.add_heading('NOTIFICATION_SYSTEM.md (+50 lineas)', level=3)
    doc.add_paragraph('Documentacion basica del sistema de notificaciones.')

    # Archivos modificados
    doc.add_heading('2. Archivos Modificados', level=2)

    add_table(doc,
        ['Archivo', 'Cambio', 'Lineas Afectadas'],
        [
            ['routes.ts', 'Agrega ruta /notificaciones con lazy loading y AuthGuard', '263-271'],
            ['maestro.service.ts', 'Agrega getCompanyNotificationPreferences() y saveCompanyNotificationPreferences()', '428-434'],
            ['sidebar.component.html', 'Agrega item "Notificaciones" con icono fa-bell', 'Bloque nuevo en sidebar'],
            ['sidebar.component.scss', 'Estilos para el nuevo item de notificaciones', '+13 lineas'],
            ['sidebar.component.ts', 'Flag para mostrar/ocultar seccion notificaciones', '+3 lineas'],
            ['notification-manager.service.ts', 'Mejoras en listeners Firebase y procesamiento', '+42 lineas'],
            ['notification-preferences.service.ts', 'Ajustes en carga de preferencias del usuario', '+10 lineas'],
            ['notification.config.ts', 'Actualizacion de configuracion de canales', '+14 lineas'],
            ['ventas.service.ts', 'Agrega clienteEmail al payload de notificaciones (3 metodos)', 'Lineas ~473, ~572, ~620'],
            ['notification.component.ts', 'Ajustes menores', '+6 lineas'],
            ['app-routing.module.ts', 'Elimina ruta de test notification-test', 'Lineas 87-91 eliminadas'],
            ['es.json', 'Traducciones: Notificaciones, Mensaje de texto, Proximamente', '+12 lineas'],
        ]
    )

    # Archivos eliminados
    doc.add_heading('3. Archivos Eliminados', level=2)
    add_table(doc,
        ['Archivo', 'Lineas', 'Razon'],
        [
            ['notification-test.component.ts', '-294', 'Componente de pruebas, ya no necesario'],
            ['notification-test.module.ts', '-17', 'Modulo del componente de test eliminado'],
            ['notification-analytics.service.ts', '-652', 'Servicio de analiticas no usado en produccion'],
        ]
    )

    # Archivos excluidos
    doc.add_heading('4. Archivos Excluidos del Merge', level=2)
    add_table(doc,
        ['Archivo', 'Razon'],
        [
            ['src/environments/environment.ts', 'Se conservo el de la rama destino (urlApi configurado)'],
            ['.claude/settings.local.json', 'Archivo local de configuracion'],
        ]
    )

    doc.add_page_break()

    # PARTE 2: BACKEND
    doc.add_heading('PARTE 2: Backend (katuq_admin_back_firebase)', level=1)
    doc.add_paragraph('Commit: a46c21a - Notification_SMS_Back_1')
    doc.add_paragraph('Resumen: 5 archivos modificados, +400 lineas, -126 lineas')

    # orders.js
    doc.add_heading('1. controllers/orders.js (+361/-45 lineas)', level=2)

    doc.add_heading('1.1 Constantes y funciones nuevas', level=3)
    add_table(doc,
        ['Elemento', 'Linea', 'Descripcion'],
        [
            ['STORAGE_BUCKET', '32', 'process.env.FIREBASE_STORAGE_BUCKET || "julsmind-katuq.appspot.com"'],
            ['DEFAULT_NOTIFICATION_PREFS', '33-36', 'Objeto con 6 claves de preferencia, todas false por defecto'],
            ['loadCompanyNotificationPrefs(company, channel)', '44-55', 'Carga preferencias desde Firestore. Reemplaza codigo duplicado en 3 funciones'],
            ['saveEmailPreview(html, orderData, templateType)', '62-85', 'Guarda HTML en Storage, metadata en Firestore, retorna URL corta o Storage'],
            ['sendDirectSms(orderData, previousOrderData, previewUrl)', '3269-3354', 'SMS para cambios de estado. Verifica env vars, preferencias, mapea estados a templates'],
            ['sendDirectCreatedSms(orderData, previewUrl)', '3360-3411', 'SMS para pedido recien creado (ORDER_CREATED)'],
        ]
    )

    doc.add_heading('1.2 Funciones modificadas', level=3)
    add_table(doc,
        ['Funcion', 'Lineas', 'Cambio'],
        [
            ['sendDirectTemplateEmail()', '3042, 3149-3155', 'Usa loadCompanyNotificationPrefs(). Agrega saveEmailPreview(). Retorna previewUrl'],
            ['sendDirectCreatedEmail()', '3257', 'Agrega saveEmailPreview(). Retorna previewUrl'],
            ['edit()', '3430-3445', 'Captura previewUrl de email, lo pasa a sendDirectSms()'],
            ['editByTransporter()', '3614-3627', 'Mismo cambio que edit(): agrega sendDirectSms()'],
            ['create()', '4503-4525', 'Agrega saveEmailPreview() del email legacy + sendDirectCreatedSms()'],
            ['sendEmail()', '6788-6802', 'Agrega LEGACY SKIP para ALLOWED_NOTIFICATION_COMPANIES'],
        ]
    )

    # notificationPreferences.js
    doc.add_heading('2. controllers/notificationPreferences.js (+137/-126 lineas)', level=2)
    doc.add_paragraph('Cambio: Reordenacion de rutas para fix de bug critico.')
    doc.add_paragraph('Bug: Las rutas genericas /:userId interceptaban /company/:companyName porque Express evalua en orden.')

    add_table(doc,
        ['Orden', 'Antes (roto)', 'Despues (corregido)'],
        [
            ['1', 'GET /:userId (linea 10)', 'GET /company/:companyName (linea 14)'],
            ['2', 'PUT /:userId (linea 38)', 'PUT /company/:companyName (linea 52)'],
            ['3', 'PUT /:userId/types/:type (linea 59)', 'GET /:userId (linea 78)'],
            ['4', 'GET /company/:companyName (linea 84)', 'PUT /:userId (linea 105)'],
            ['5', 'PUT /company/:companyName (linea 122)', 'PUT /:userId/types/:type (linea 126)'],
        ]
    )

    # index.js
    doc.add_heading('3. functions/index.js (+14 lineas)', level=2)
    doc.add_paragraph('Lineas 540-552: Nuevo endpoint GET /v1/e/:id')
    doc.add_paragraph('Busca en Firestore collection email_previews por document ID. Si existe y tiene storageUrl, hace redirect 301. Si no existe, retorna 404 con HTML amigable.')

    # notifications.js
    doc.add_heading('4. routers/notifications.js (-29 lineas)', level=2)
    doc.add_paragraph('Lineas eliminadas: 144-171')
    doc.add_paragraph('Se elimino el endpoint GET /email-view/:id porque duplicaba la funcionalidad de /v1/e/:id en index.js. Tambien se elimino el anti-pattern de require("firebase-admin") dentro del handler.')

    # smsTemplates.js
    doc.add_heading('5. services/notifications/smsTemplates.js (+6/-6 lineas)', level=2)
    doc.add_paragraph('Se agrego {{linkPedido}} a todos los templates SMS:')
    add_table(doc,
        ['Template', 'Linea', 'Cambio'],
        [
            ['ORDER_CREATED', '10', 'Reemplaza texto final por "Detalles: {{linkPedido}}"'],
            ['PAYMENT_Aprobado', '14', 'Reemplaza texto final por "Detalles: {{linkPedido}}"'],
            ['ORDER_ProducidoTotalmente', '18', 'Agrega "Detalles: {{linkPedido}}"'],
            ['ORDER_Despachado', '22', 'Reemplaza Transportador por Guia + link'],
            ['ORDER_Entregado', '26', 'Reemplaza texto final por "Detalles: {{linkPedido}}"'],
            ['ORDER_Rechazado', '30', 'Reemplaza texto final por "Detalles: {{linkPedido}}"'],
        ]
    )

    # notificationQueue.js
    doc.add_heading('6. services/notifications/notificationQueue.js (+1/-1 linea)', level=2)
    doc.add_paragraph('Linea 267: Actualiza PRODUCTION_CUTOFF de 2026-03-16T23:25:00Z a 2026-03-18T23:59:00Z')

    doc.add_page_break()

    # PARTE 3: CONFIGURACION .env
    doc.add_heading('PARTE 3: Configuracion .env', level=1)

    doc.add_heading('Variables Agregadas', level=2)
    add_table(doc,
        ['Variable', 'Valor', 'Descripcion'],
        [
            ['BACKEND_PUBLIC_URL', 'https://back.katuq.com', 'URL base para generar URLs cortas en SMS'],
            ['FIREBASE_STORAGE_BUCKET', '(usa default)', 'Variable opcional, fallback a julsmind-katuq.appspot.com'],
        ]
    )

    doc.add_heading('Variables Existentes Relevantes', level=2)
    add_table(doc,
        ['Variable', 'Valor Actual', 'Para Produccion', 'Descripcion'],
        [
            ['NOTIFICATIONS_PUBLIC_MODE', 'false', 'Cambiar a true', 'Modo test vs produccion'],
            ['ENABLE_SMS_NOTIFICATIONS', 'true', 'true', 'SMS habilitado'],
            ['ENABLE_EMAIL_NOTIFICATIONS', 'true', 'true', 'Email habilitado'],
            ['ALLOWED_NOTIFICATION_COMPANIES', 'Tienda Demo KAI Import', 'Vaciar para todas', 'Filtro de empresas'],
            ['LABSMOBILE_USERNAME', 'gerencia@almara.com.co', 'Verificar', 'Usuario labsMobile'],
            ['LABSMOBILE_API_KEY', '2kPK...YRJP', 'Verificar', 'API Key labsMobile'],
            ['LABSMOBILE_SENDER', 'Katuq', 'Katuq', 'Remitente SMS'],
            ['NOTIFICATIONS_TEST_EMAIL', 'santygarcia...@gmail.com', 'N/A (solo test)', 'Email pruebas'],
            ['NOTIFICATIONS_TEST_PHONE', '573024218994', 'N/A (solo test)', 'Telefono pruebas'],
        ]
    )

    doc.add_heading('Checklist para Deploy a Produccion', level=2)
    checklist = [
        'Cambiar NOTIFICATIONS_PUBLIC_MODE=true en el servidor',
        'Agregar BACKEND_PUBLIC_URL=https://back.katuq.com en el servidor',
        'Verificar que LABSMOBILE_USERNAME y LABSMOBILE_API_KEY son correctas',
        'Decidir si quitar el filtro ALLOWED_NOTIFICATION_COMPANIES (vaciar = todas las empresas)',
        'Verificar saldo de creditos en labsMobile',
    ]
    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')

    filepath = os.path.join(OUTPUT_DIR, 'NOTIFICACIONES_MERGE_DETALLE.docx')
    doc.save(filepath)
    print(f'Documento de merge generado: {filepath}')


if __name__ == '__main__':
    generate_technical_doc()
    generate_merge_doc()
    print('Documentos generados exitosamente.')
